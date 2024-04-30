import pandas as pd
import json
from docx import Document

"""
一个转换类，将主流文件互相转换
"""


class FileConverter:
    def __init__(self, source_file_path):
        self.source_file_path = source_file_path

    def csv_to_xlsx(self):
        df = pd.read_csv(self.source_file_path)
        new_file_path = self.source_file_path.with_suffix(".xlsx")
        df.to_excel(new_file_path, index=False)
        print(f"转换完成: {self} -> {new_file_path}")

    def xlsx_to_csv(self):
        df = pd.read_excel(self.source_file_path)
        new_file_path = self.source_file_path.with_suffix(".csv")
        df.to_csv(new_file_path, index=False)
        print(f"转换完成: {self.source_file_path} -> {new_file_path}")

    def word_to_json(self):
        document = Document(self.source_file_path)
        data = []
        for paragraph in document.paragraphs:
            data.append({"text": paragraph.text})
        # ... 处理 JSON 数据 ...
        print(f"转换完成: {self.source_file_path} -> JSON 数据")

    def word_to_txt(self):
        document = Document(self.source_file_path)
        text = "\n".join([paragraph.text for paragraph in document.paragraphs])
        new_file_path = self.source_file_path.with_suffix(".txt")
        with open(new_file_path, "w", encoding='utf-8') as f:
            f.write(text)
        print(f"转换完成: {self.source_file_path} -> {new_file_path}")

    def str_to_json(self):
        pass


if __name__ == '__main__':
    path = ''
    translate_file_type = FileConverter(path)
    # translate_file_type.source_file_path = r'C:\Users\wuxd-a\PycharmProjects\pythonProject6\4.25'
